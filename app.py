import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
import datetime

# ── scikit-learn ──────────────────────────────────────────────────────────────
from sklearn.linear_model import LinearRegression, LogisticRegression
from sklearn.neural_network import MLPRegressor, MLPClassifier
from sklearn.ensemble import RandomForestRegressor, RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.metrics import (
    mean_squared_error, accuracy_score, f1_score,
    confusion_matrix, ConfusionMatrixDisplay,
)
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="ML Workbench",
    page_icon="⚗️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
html, body, [data-testid="stAppViewContainer"],
[data-testid="stSidebar"], .main {
    background-color: #f7f8fc !important;
    color: #1a1a2e !important;
}
[data-testid="stSidebar"] {
    background-color: #ffffff !important;
    border-right: 1px solid #e2e5ef;
    min-width: 230px !important;
    max-width: 230px !important;
}
.step-badge {
    display:flex; align-items:center; gap:10px; padding:10px 14px;
    border-radius:8px; margin-bottom:6px; font-size:.88rem; font-weight:500;
}
.step-badge.active { background:#e8f0fe; color:#1a56db; border-left:3px solid #1a56db; }
.step-badge.done   { background:#f0fdf4; color:#15803d; border-left:3px solid #22c55e; }
.step-badge.locked { background:#f1f5f9; color:#94a3b8; border-left:3px solid #cbd5e1; }
.card {
    background:#ffffff; border:1px solid #e2e5ef; border-radius:12px;
    padding:24px 28px; margin-bottom:20px; box-shadow:0 1px 4px rgba(0,0,0,.04);
}
.dim-pill {
    display:inline-block; background:#e8f0fe; color:#1a56db;
    border-radius:20px; padding:4px 16px; font-size:.95rem; font-weight:600; margin:2px 4px;
}
div.stButton > button {
    background:#1a56db; color:white; border:none; border-radius:8px;
    padding:.45rem 1.4rem; font-weight:600; letter-spacing:.03em;
}
div.stButton > button:hover { background:#1648c0; }
[data-testid="stMetric"] {
    background:#f8faff; border:1px solid #dde3f0; border-radius:10px; padding:12px 16px;
}
details { border:1px solid #e2e5ef !important; border-radius:8px !important; }
.stAlert { border-radius:8px !important; }
h2 { font-size:1.25rem !important; font-weight:700 !important; color:#1a1a2e !important; }
h3 { font-size:1.05rem !important; font-weight:600 !important; color:#374151 !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════
defaults = {
    "step": 1,
    "raw_df": None,
    "confirmed_df": None,
    "type_map": {},
    "step1_done": False,
    "step2_done": False,
    "step3_done": False,
    "model_results": {},
    "last_qoi": None,
    "last_task": None,
    "last_train_pct": 80,
    "file_name": "",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def infer_type(series):
    return "Numerical" if pd.api.types.is_numeric_dtype(series) else "Categorical"

def apply_types(df, type_map):
    df = df.copy()
    for col, typ in type_map.items():
        if typ == "Numerical":
            df[col] = pd.to_numeric(df[col], errors="coerce")
        else:
            df[col] = df[col].astype(str)
    return df

def encode_for_model(df, type_map, qoi):
    df = df.copy()
    cat_cols = [c for c, t in type_map.items() if t == "Categorical" and c != qoi]
    df = pd.get_dummies(df, columns=cat_cols, drop_first=True)
    return df

def sidebar_steps():
    with st.sidebar:
        st.markdown("### ⚗️ ML Workbench")
        st.markdown("---")
        steps = [(1,"Load Data"),(2,"Variable Types"),(3,"Model & Compare"),(4,"Report")]
        for num, label in steps:
            if num == st.session_state.step:
                cls = "active"
            elif num < st.session_state.step:
                cls = "done"
            else:
                cls = "locked"
            icon = {"active":"▶","done":"✓","locked":"🔒"}[cls]
            st.markdown(
                f'<div class="step-badge {cls}">{icon} &nbsp; Step {num} — {label}</div>',
                unsafe_allow_html=True,
            )
        st.markdown("---")
        st.caption("Steps unlock sequentially.")

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 1 — LOAD DATA
# ══════════════════════════════════════════════════════════════════════════════
def step1():
    st.markdown("## Step 1 — Load Data")
    st.markdown('<div class="card">', unsafe_allow_html=True)
    uploaded = st.file_uploader("Upload an Excel file (.xlsx, single sheet)", type=["xlsx"])
    if uploaded:
        try:
            df = pd.read_excel(BytesIO(uploaded.read()), sheet_name=0)
            st.session_state.raw_df = df
            st.session_state.file_name = uploaded.name
            st.success(f"**{uploaded.name}** loaded successfully.")
            c1, c2 = st.columns(2)
            c1.markdown(f'<div style="margin-top:10px">Rows &nbsp;<span class="dim-pill">{df.shape[0]:,}</span></div>', unsafe_allow_html=True)
            c2.markdown(f'<div style="margin-top:10px">Columns &nbsp;<span class="dim-pill">{df.shape[1]}</span></div>', unsafe_allow_html=True)
            st.markdown("**Preview (first 5 rows):**")
            st.dataframe(df.head(), use_container_width=True)
            if st.button("Confirm & Proceed to Step 2"):
                st.session_state.step1_done = True
                st.session_state.step = 2
                st.session_state.type_map = {col: infer_type(df[col]) for col in df.columns}
                st.rerun()
        except Exception as e:
            st.error(f"Could not read file: {e}")
    st.markdown("</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 2 — VARIABLE TYPES
# ══════════════════════════════════════════════════════════════════════════════
def step2():
    st.markdown("## Step 2 — Variable Types")
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown("Review inferred types. Override where needed, then confirm.")
    df       = st.session_state.raw_df
    type_map = st.session_state.type_map.copy()
    hc = st.columns([3,2,2])
    hc[0].markdown("**Variable**")
    hc[1].markdown("**Inferred Type**")
    hc[2].markdown("**Override**")
    st.markdown('<hr style="margin:4px 0 10px 0;border-color:#e2e5ef">', unsafe_allow_html=True)
    new_type_map = {}
    for col in df.columns:
        row      = st.columns([3,2,2])
        inferred = type_map[col]
        row[0].markdown(f"`{col}`")
        row[1].markdown(
            f'<span style="color:{"#1a56db" if inferred=="Numerical" else "#7c3aed"};font-weight:600">{inferred}</span>',
            unsafe_allow_html=True,
        )
        opts = ["Numerical","Categorical"]
        new_type_map[col] = row[2].selectbox(
            "", opts, index=opts.index(inferred), key=f"type_{col}", label_visibility="collapsed"
        )
    st.markdown("</div>", unsafe_allow_html=True)
    if st.button("Confirm Types & Proceed to Step 3"):
        st.session_state.type_map    = new_type_map
        st.session_state.confirmed_df = apply_types(df, new_type_map)
        st.session_state.step2_done  = True
        st.session_state.step        = 3
        st.rerun()
    if st.button("Back to Step 1"):
        st.session_state.step = 1
        st.session_state.step1_done = False
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 3 — MODEL & COMPARE
# ══════════════════════════════════════════════════════════════════════════════
def step3():
    st.markdown("## Step 3 — Define QOI & Build Models")
    df       = st.session_state.confirmed_df
    type_map = st.session_state.type_map

    # QOI
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown("### Target Variable (QOI)")
    eligible = [c for c in df.columns
                if type_map[c]=="Numerical" or
                   (type_map[c]=="Categorical" and df[c].nunique()==2)]
    if not eligible:
        st.error("No eligible QOI columns (need numerical or binary categorical).")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    qoi  = st.selectbox("Select QOI", eligible,
                        help="Non-binary categoricals excluded. Numerical->regression; Binary categorical->classification.")
    task = "regression" if type_map[qoi]=="Numerical" else "classification"
    st.session_state.last_qoi  = qoi
    st.session_state.last_task = task

    st.info(f"**Task:** {'Regression' if task=='regression' else 'Binary Classification'}  \n"
            f"**Models:** {'OLS · Neural Network · Random Forest' if task=='regression' else 'Logistic Regression · Neural Network · Random Forest'}")

    train_pct = st.slider("Training set size (%)", 50, 95, 80, 5)
    st.session_state.last_train_pct = train_pct
    st.markdown("</div>", unsafe_allow_html=True)

    # Prepare data
    df_enc = encode_for_model(df, type_map, qoi).dropna(subset=[qoi])
    le = None
    if task == "classification":
        le = LabelEncoder()
        df_enc[qoi] = le.fit_transform(df_enc[qoi].astype(str))

    feature_cols = [c for c in df_enc.columns if c != qoi]
    X = df_enc[feature_cols].copy()
    y = df_enc[qoi].copy()
    mask = X.notna().all(axis=1)
    X, y = X[mask], y[mask]
    scaler   = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    X_tr, X_te, y_tr, y_te = train_test_split(X_scaled, y, train_size=train_pct/100, random_state=42)

    model_results = st.session_state.model_results

    def run_model(name, mdl):
        mdl.fit(X_tr, y_tr)
        if task == "regression":
            tr_p = mdl.predict(X_tr)
            te_p = mdl.predict(X_te)
            return {"name": name,
                    "train_mse":  mean_squared_error(y_tr, tr_p),
                    "test_mse":   mean_squared_error(y_te, te_p),
                    "train_rmse": np.sqrt(mean_squared_error(y_tr, tr_p)),
                    "test_rmse":  np.sqrt(mean_squared_error(y_te, te_p))}
        else:
            tp, ep = mdl.predict(X_tr), mdl.predict(X_te)
            return {"name": name,
                    "train_acc": accuracy_score(y_tr, tp),
                    "test_acc":  accuracy_score(y_te, ep),
                    "train_f1":  f1_score(y_tr, tp, zero_division=0),
                    "test_f1":   f1_score(y_te, ep, zero_division=0),
                    "cm_train":  confusion_matrix(y_tr, tp),
                    "cm_test":   confusion_matrix(y_te, ep),
                    "le": le}

    # ── Model 1 ───────────────────────────────────────────────────────────────
    m1_label = "OLS Linear Regression" if task=="regression" else "Logistic Regression (GLM)"
    with st.expander(f"**Model 1 — {m1_label}**", expanded=True):
        m1_params = {}
        if task == "regression":
            st.markdown("*OLS fits analytically — no tunable hyperparameters.*")
        else:
            ca, cb = st.columns(2)
            m1_params["C"] = ca.number_input(
                "Regularisation C", 0.001, 100.0, 1.0, 0.1,
                help="Inverse regularisation strength. Smaller = stronger. Range: 0.001–100.", key="m1_C")
            m1_params["max_iter"] = cb.number_input(
                "Max iterations", 100, 10000, 1000, 100,
                help="Maximum solver iterations. Range: 100–10 000.", key="m1_iter")
        if st.button(f"Run {m1_label}", key="run_m1"):
            with st.spinner("Fitting..."):
                mdl = LinearRegression() if task=="regression" else \
                      LogisticRegression(C=m1_params["C"], max_iter=int(m1_params["max_iter"]), random_state=42)
                model_results["m1"] = run_model(m1_label, mdl)
                model_results["m1"]["params"] = {} if task=="regression" else m1_params
            st.success("Done!")
        if "m1" in model_results:
            _show_metrics(model_results["m1"], task)

    # ── Model 2 ───────────────────────────────────────────────────────────────
    with st.expander("**Model 2 — Neural Network (MLP)**", expanded=True):
        ca, cb, cc = st.columns(3)
        m2_layers = ca.text_input("Hidden layer sizes (comma-separated)", "100,50",
            help="E.g. '100,50' = two hidden layers. Each value: 1–1 000.", key="m2_layers")
        m2_alpha  = cb.number_input("L2 regularisation (alpha)", 1e-6, 10.0, 1e-4, format="%.6f",
            help="L2 penalty. Range: 0.000001–10.", key="m2_alpha")
        m2_iter   = cc.number_input("Max iterations", 50, 2000, 500, 50,
            help="Training epochs. Range: 50–2 000.", key="m2_iter")
        if st.button("Run Neural Network", key="run_m2"):
            try:
                layers = tuple(int(x.strip()) for x in m2_layers.split(",") if x.strip())
                assert all(1 <= n <= 1000 for n in layers), "Each layer size must be 1–1 000."
                with st.spinner("Fitting..."):
                    MLP = MLPRegressor if task=="regression" else MLPClassifier
                    mdl = MLP(hidden_layer_sizes=layers, alpha=m2_alpha, max_iter=int(m2_iter), random_state=42)
                    model_results["m2"] = run_model("Neural Network (MLP)", mdl)
                    model_results["m2"]["params"] = {"hidden_layers": m2_layers,
                                                      "alpha": m2_alpha,
                                                      "max_iter": int(m2_iter)}
                st.success("Done!")
            except Exception as e:
                st.error(f"Invalid parameters: {e}")
        if "m2" in model_results:
            _show_metrics(model_results["m2"], task)

    # ── Model 3 ───────────────────────────────────────────────────────────────
    with st.expander("**Model 3 — Random Forest**", expanded=True):
        ca, cb, cc = st.columns(3)
        m3_n     = ca.number_input("Number of trees", 10, 1000, 100, 10,
            help="More trees = more stable, slower. Range: 10–1 000.", key="m3_n")
        m3_depth = cb.number_input("Max depth (0 = unlimited)", 0, 100, 0, 1,
            help="Max depth per tree. 0 = unlimited. Range: 0–100.", key="m3_depth")
        m3_split = cc.number_input("Min samples to split", 2, 100, 2, 1,
            help="Min samples to split an internal node. Range: 2–100.", key="m3_split")
        if st.button("Run Random Forest", key="run_m3"):
            with st.spinner("Fitting..."):
                depth = None if m3_depth == 0 else int(m3_depth)
                RF  = RandomForestRegressor if task=="regression" else RandomForestClassifier
                mdl = RF(n_estimators=int(m3_n), max_depth=depth,
                         min_samples_split=int(m3_split), random_state=42)
                model_results["m3"] = run_model("Random Forest", mdl)
                model_results["m3"]["params"] = {
                    "n_estimators":      int(m3_n),
                    "max_depth":         "unlimited" if m3_depth == 0 else int(m3_depth),
                    "min_samples_split": int(m3_split),
                }
            st.success("Done!")
        if "m3" in model_results:
            _show_metrics(model_results["m3"], task)

    # ── Summary ───────────────────────────────────────────────────────────────
    run_keys = [k for k in ["m1","m2","m3"] if k in model_results]
    if run_keys:
        st.markdown("---")
        st.markdown("### Combined Model Comparison")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        _show_summary(model_results, run_keys, task)
        st.markdown("</div>", unsafe_allow_html=True)

        if len(run_keys) == 3:
            st.session_state.step3_done = True
            st.markdown("---")
            if st.button("Proceed to Report (Step 4)"):
                st.session_state.step = 4
                st.rerun()

    if st.button("Back to Step 2"):
        st.session_state.step        = 2
        st.session_state.step2_done  = False
        st.session_state.model_results = {}
        st.rerun()


def _show_metrics(r, task):
    if task == "regression":
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Train MSE",  f"{r['train_mse']:.4f}")
        c2.metric("Test MSE",   f"{r['test_mse']:.4f}")
        c3.metric("Train RMSE", f"{r['train_rmse']:.4f}")
        c4.metric("Test RMSE",  f"{r['test_rmse']:.4f}")
    else:
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Train Accuracy", f"{r['train_acc']:.4f}")
        c2.metric("Test Accuracy",  f"{r['test_acc']:.4f}")
        c3.metric("Train F1",       f"{r['train_f1']:.4f}")
        c4.metric("Test F1",        f"{r['test_f1']:.4f}")


def _show_summary(model_results, run_keys, task):
    if task == "regression":
        rows = [{"Model":      model_results[k]["name"],
                 "Train MSE":  round(model_results[k]["train_mse"],  6),
                 "Test MSE":   round(model_results[k]["test_mse"],   6),
                 "Train RMSE": round(model_results[k]["train_rmse"], 6),
                 "Test RMSE":  round(model_results[k]["test_rmse"],  6)} for k in run_keys]
        summary = pd.DataFrame(rows).set_index("Model")
        def hl_min(s):
            return ["background-color:#dcfce7;color:#15803d;font-weight:700"
                    if v == s.min() else "" for v in s]
        st.dataframe(summary.style.apply(hl_min, subset=["Test MSE","Test RMSE"]),
                     use_container_width=True)
        st.caption("Green = best (lowest) value.")
    else:
        rows = [{"Model":           model_results[k]["name"],
                 "Train Accuracy":  round(model_results[k]["train_acc"], 4),
                 "Test Accuracy":   round(model_results[k]["test_acc"],  4),
                 "Train F1":        round(model_results[k]["train_f1"],  4),
                 "Test F1":         round(model_results[k]["test_f1"],   4)} for k in run_keys]
        summary = pd.DataFrame(rows).set_index("Model")
        def hl_max(s):
            return ["background-color:#dcfce7;color:#15803d;font-weight:700"
                    if v == s.max() else "" for v in s]
        st.dataframe(summary.style.apply(hl_max, subset=["Test Accuracy","Test F1"]),
                     use_container_width=True)
        st.caption("Green = best (highest) value.")
        st.markdown("**Confusion Matrices (Test Set)**")
        cm_cols = st.columns(len(run_keys))
        for i, k in enumerate(run_keys):
            r   = model_results[k]
            fig, ax = plt.subplots(figsize=(3.2, 2.8))
            fig.patch.set_facecolor("#ffffff"); ax.set_facecolor("#ffffff")
            ConfusionMatrixDisplay(r["cm_test"], display_labels=r["le"].classes_).plot(
                ax=ax, colorbar=False, cmap="Blues")
            ax.set_title(r["name"], fontsize=9, fontweight="bold", color="#1a1a2e")
            ax.tick_params(labelsize=8); plt.tight_layout()
            cm_cols[i].pyplot(fig); plt.close(fig)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX BUILDER  (python-docx, generated at runtime)
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, hex_color="CCCCCC", size=4):
    tcPr    = cell._tc.get_or_add_tcPr()
    tcBords = OxmlElement("w:tcBorders")
    for side in ("top","bottom","left","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        tcBords.append(el)
    tcPr.append(tcBords)


def _add_para(doc, text, bold=False, italic=False, size=11,
              rgb=None, space_before=0, space_after=6,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    run.font.name   = "Calibri"
    if rgb:
        run.font.color.rgb = RGBColor(*[int(rgb[i:i+2], 16) for i in (0,2,4)])
    return p


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.name = "Calibri"
    if level == 1:
        run.font.size      = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        # underline rule
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single"); bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1");      bot.set(qn("w:color"), "1A56DB")
        pBdr.append(bot); pPr.append(pBdr)
    else:
        run.font.size      = Pt(12)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def _add_kv(doc, key, value):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{key}: ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r1.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    r2 = p.add_run(str(value))
    r2.font.size = Pt(10); r2.font.name = "Calibri"
    r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_bullet(doc, text):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


# ── model static metadata ────────────────────────────────────────────────────
MODEL_META = {
    "regression": {
        "m1": {
            "title": "Model 1 — Ordinary Least Squares (OLS) Linear Regression",
            "description": (
                "OLS minimises the residual sum of squares between observed and predicted values. "
                "Coefficients are estimated analytically via the normal equations (or QR decomposition), "
                "making it deterministic and free of tunable hyperparameters. "
                "OLS assumes a linear relationship between predictors and the response, "
                "homoscedastic and uncorrelated errors, and the absence of perfect multicollinearity."
            ),
            "assumptions": [
                "Linearity: E[Y|X] is linear in X.",
                "Independence of error terms.",
                "Homoscedasticity: constant error variance across all fitted values.",
                "No perfect multicollinearity among predictor variables.",
            ],
        },
        "m2": {
            "title": "Model 2 — Multi-Layer Perceptron Regressor (Neural Network)",
            "description": (
                "A fully-connected feed-forward neural network with one or more hidden layers. "
                "Each neuron applies a non-linear activation function (ReLU by default in sklearn). "
                "Weights are optimised by back-propagation using the Adam solver, minimising MSE. "
                "L2 regularisation (alpha) penalises large weights to reduce overfitting. "
                "Features must be on a comparable scale — standardisation is applied before fitting."
            ),
            "assumptions": [
                "Sufficient data relative to the number of parameters (depth x width).",
                "Features are standardised before training (applied in this pipeline).",
                "Convergence within max_iter epochs is not guaranteed for all datasets.",
            ],
        },
        "m3": {
            "title": "Model 3 — Random Forest Regressor",
            "description": (
                "An ensemble of decision trees, each trained on a bootstrap sample of the data "
                "with a random subset of features considered at each split (bagging + feature randomness). "
                "Predictions are the mean of all individual tree predictions. "
                "Random Forests are robust to outliers and naturally capture non-linear relationships "
                "without requiring feature scaling."
            ),
            "assumptions": [
                "Sufficient number of trees to reduce ensemble variance (typically >= 100).",
                "max_depth controls individual tree complexity; unlimited depth may overfit on small datasets.",
                "Feature scaling is not required.",
            ],
        },
    },
    "classification": {
        "m1": {
            "title": "Model 1 — Binary Logistic Regression (GLM)",
            "description": (
                "A generalised linear model that models the log-odds of the binary outcome as a linear "
                "combination of predictors. The sigmoid function maps linear predictions to probabilities in [0,1]. "
                "Coefficients are estimated by maximum likelihood. "
                "The regularisation parameter C (inverse of lambda) controls the trade-off between "
                "fit quality and coefficient magnitude."
            ),
            "assumptions": [
                "Binary outcome variable.",
                "Log-odds are linearly related to the predictor variables.",
                "Independence of observations.",
                "Absence of severe multicollinearity among predictors.",
            ],
        },
        "m2": {
            "title": "Model 2 — Multi-Layer Perceptron Classifier (Neural Network)",
            "description": (
                "A feed-forward neural network for binary classification, using a logistic output layer. "
                "Trained with back-propagation (Adam solver) minimising cross-entropy loss. "
                "L2 regularisation (alpha) helps mitigate overfitting. "
                "Features are standardised before training."
            ),
            "assumptions": [
                "Sufficient data for the chosen architecture.",
                "Features standardised before training (applied in this pipeline).",
                "Convergence within max_iter epochs is not guaranteed.",
            ],
        },
        "m3": {
            "title": "Model 3 — Random Forest Classifier",
            "description": (
                "An ensemble of classification trees, each trained on a bootstrap sample with random "
                "feature subsets considered at each split. Final prediction is by majority vote. "
                "Captures complex non-linear decision boundaries and is robust to irrelevant features. "
                "Feature scaling is not required."
            ),
            "assumptions": [
                "Sufficient trees to reduce ensemble variance.",
                "max_depth limits per-tree complexity to control overfitting.",
                "Feature scaling is not required.",
            ],
        },
    },
}


def build_docx_report(model_results, task, qoi, train_pct, file_name):
    doc = DocxDocument()

    # Page size: US Letter, 1-inch margins (twips: 1440 per inch)
    sec = doc.sections[0]
    sec.page_width  = 12240
    sec.page_height = 15840
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec, attr, 1440)

    # ── Title ─────────────────────────────────────────────────────────────────
    _add_para(doc, "ML Workbench — Model Comparison Report",
              bold=True, size=18, rgb="1A1A2E", space_after=2,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Generated: {datetime.datetime.now().strftime('%d %B %Y, %H:%M')}",
              size=10, rgb="6B7280", space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Source file: {file_name}",
              size=10, rgb="6B7280", space_after=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 1. Problem Setup ──────────────────────────────────────────────────────
    _add_heading(doc, "1. Problem Setup")
    _add_kv(doc, "Quantity of Interest (QOI)", qoi)
    _add_kv(doc, "Task type", "Regression" if task=="regression" else "Binary Classification")
    _add_kv(doc, "Training / test split", f"{train_pct}% / {100 - train_pct}%")
    _add_kv(doc, "Feature encoding",
            "One-hot encoding for categorical predictors; all features standardised (zero mean, unit variance)")
    doc.add_paragraph()

    # ── 2. Model Descriptions ─────────────────────────────────────────────────
    _add_heading(doc, "2. Model Descriptions & Parameters")
    run_keys = [k for k in ["m1","m2","m3"] if k in model_results]

    for k in run_keys:
        r    = model_results[k]
        meta = MODEL_META[task][k]

        _add_heading(doc, meta["title"], level=2)
        _add_para(doc, meta["description"], size=10, space_after=4)

        _add_para(doc, "Key assumptions:", bold=True, size=10, space_after=2)
        for assumption in meta["assumptions"]:
            _add_bullet(doc, assumption)

        params = r.get("params", {})
        if params:
            _add_para(doc, "Parameters used:", bold=True, size=10, space_before=6, space_after=2)
            for pk, pv in params.items():
                _add_kv(doc, pk, pv)
        else:
            _add_para(doc, "Parameters: Fitted analytically — none required.",
                      italic=True, size=10, space_after=4)
        doc.add_paragraph()

    # ── 3. Performance Comparison Table ───────────────────────────────────────
    _add_heading(doc, "3. Performance Comparison")

    if task == "regression":
        _add_para(doc,
                  "The table below reports Mean Squared Error (MSE) and Root Mean Squared Error (RMSE) "
                  "on both training and test sets. Lower values indicate better fit. "
                  "Test-set performance is the primary indicator of generalisation.",
                  size=10, space_after=8)
        col_labels = ["Model", "Train MSE", "Test MSE", "Train RMSE", "Test RMSE"]
        col_widths  = [2520, 1710, 1710, 1710, 1710]   # sum = 9360 (US Letter body)
        rows_data   = [[r["name"],
                        f"{r['train_mse']:.6f}", f"{r['test_mse']:.6f}",
                        f"{r['train_rmse']:.6f}", f"{r['test_rmse']:.6f}"]
                       for r in [model_results[k] for k in run_keys]]
        best_idx = int(np.argmin([model_results[k]["test_mse"] for k in run_keys]))
        best_note = "Green row = lowest Test MSE (best generalisation)."
    else:
        _add_para(doc,
                  "The table below reports Accuracy and F1 Score on training and test sets. "
                  "Higher values indicate better classification performance. "
                  "Confusion matrices appear in Section 4.",
                  size=10, space_after=8)
        col_labels = ["Model", "Train Acc.", "Test Acc.", "Train F1", "Test F1"]
        col_widths  = [2520, 1710, 1710, 1710, 1710]
        rows_data   = [[r["name"],
                        f"{r['train_acc']:.4f}", f"{r['test_acc']:.4f}",
                        f"{r['train_f1']:.4f}",  f"{r['test_f1']:.4f}"]
                       for r in [model_results[k] for k in run_keys]]
        best_idx = int(np.argmax([model_results[k]["test_acc"] for k in run_keys]))
        best_note = "Green row = highest Test Accuracy (best generalisation)."

    tbl = doc.add_table(rows=1 + len(run_keys), cols=len(col_labels))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    HEADER_BG = "1A56DB"
    ROW_ALT   = ["FFFFFF", "EEF2FF"]
    BEST_BG   = "DCFCE7"

    # Header
    for ci, (cell, label, w) in enumerate(zip(tbl.rows[0].cells, col_labels, col_widths)):
        cell.width = w
        _set_cell_bg(cell, HEADER_BG)
        _set_cell_borders(cell, "FFFFFF")
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(label)
        run.bold = True; run.font.size = Pt(10); run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row_vals in enumerate(rows_data):
        tbl_row = tbl.rows[ri + 1]
        bg = BEST_BG if ri == best_idx else ROW_ALT[ri % 2]
        for ci, (cell, val, w) in enumerate(zip(tbl_row.cells, row_vals, col_widths)):
            cell.width = w
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell, "CCCCCC")
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(val)
            run.bold = (ri == best_idx)
            run.font.size = Pt(10); run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    _add_para(doc, best_note, italic=True, size=9, rgb="15803D", space_before=4, space_after=12)

    # ── 4. Confusion Matrices (classification only) ───────────────────────────
    if task == "classification":
        doc.add_page_break()
        _add_heading(doc, "4. Confusion Matrices (Test Set)")
        _add_para(doc,
                  "Each matrix shows predicted vs. actual class labels on the held-out test set. "
                  "Diagonal cells represent correct predictions; off-diagonal cells are misclassifications.",
                  size=10, space_after=10)

        for k in run_keys:
            r    = model_results[k]
            meta = MODEL_META["classification"][k]
            _add_heading(doc, meta["title"], level=2)

            fig, ax = plt.subplots(figsize=(4, 3.5))
            fig.patch.set_facecolor("white"); ax.set_facecolor("white")
            ConfusionMatrixDisplay(r["cm_test"], display_labels=r["le"].classes_).plot(
                ax=ax, colorbar=False, cmap="Blues")
            ax.set_title(r["name"], fontsize=10, fontweight="bold")
            plt.tight_layout()
            img_buf = BytesIO()
            fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
            img_buf.seek(0)
            plt.close(fig)

            doc.add_picture(img_buf, width=Inches(3.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()

    # ── 5. Notes ──────────────────────────────────────────────────────────────
    sec_num = "5" if task == "classification" else "4"
    _add_heading(doc, f"{sec_num}. Notes & Interpretation Guidance")
    notes = [
        "All models were trained with a fixed random seed (42) for reproducibility.",
        "Features were standardised (zero mean, unit variance) before fitting — particularly important for OLS/Logistic Regression and MLP.",
        "Categorical predictors were one-hot encoded with drop-first to avoid perfect multicollinearity.",
        "Rows with missing values in any predictor column were excluded prior to modelling.",
        "A large gap between training and test metrics typically indicates overfitting.",
        "For the MLP, sklearn may raise a ConvergenceWarning if max_iter is too low; increase it if this occurs.",
        "Random Forest test performance stabilises as n_estimators increases; use at least 100 trees.",
    ]
    for note in notes:
        _add_bullet(doc, note)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  STEP 4 — REPORT
# ══════════════════════════════════════════════════════════════════════════════
def step4():
    st.markdown("## Step 4 — Model Report")
    st.markdown('<div class="card">', unsafe_allow_html=True)

    model_results = st.session_state.model_results
    task          = st.session_state.last_task
    qoi           = st.session_state.last_qoi
    train_pct     = st.session_state.last_train_pct
    file_name     = st.session_state.file_name
    run_keys      = [k for k in ["m1","m2","m3"] if k in model_results]

    if len(run_keys) < 3:
        st.warning("Please run all three models in Step 3 before generating the report.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    # Preview
    st.markdown("### Report Preview")
    st.markdown(
        f"**QOI:** `{qoi}`  &nbsp;|&nbsp;  "
        f"**Task:** {'Regression' if task=='regression' else 'Binary Classification'}  &nbsp;|&nbsp;  "
        f"**Split:** {train_pct} / {100-train_pct}"
    )
    st.markdown("**Models included:**")
    for k in run_keys:
        st.markdown(f"- {model_results[k]['name']}")

    st.markdown("---")
    st.markdown("**Performance summary:**")
    _show_summary(model_results, run_keys, task)

    st.markdown("---")
    st.markdown("### Generate & Download")
    st.markdown(
        "The report includes: problem setup, per-model description and parameters used, "
        "an annotated performance comparison table, confusion matrices (classification), "
        "and interpretation notes."
    )

    if st.button("Generate Report (.docx)"):
        with st.spinner("Building Word document..."):
            docx_buf = build_docx_report(model_results, task, qoi, train_pct, file_name)
        ts    = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        fname = f"ML_Report_{qoi}_{ts}.docx"
        st.download_button(
            label=f"Download {fname}",
            data=docx_buf,
            file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.success(f"Report ready — click above to download **{fname}**.")

    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("---")
    if st.button("Back to Step 3"):
        st.session_state.step = 3
        st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN ROUTER
# ══════════════════════════════════════════════════════════════════════════════
def main():
    sidebar_steps()
    step = st.session_state.step
    if step == 1:
        step1()
    elif step == 2:
        step2() if st.session_state.step1_done else st.warning("Complete Step 1 first.")
    elif step == 3:
        step3() if st.session_state.step2_done else st.warning("Complete Step 2 first.")
    elif step == 4:
        step4() if st.session_state.step3_done else st.warning("Complete Step 3 first.")


if __name__ == "__main__":
    main()
